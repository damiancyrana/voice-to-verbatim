import os
import requests
import uuid
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def translate_text(text, to_lang="pl"):
    key = ""
    endpoint = ""
    location = ""
    path = '/translate'
    constructed_url = endpoint + path

    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': [to_lang]
    }

    headers = {
        'Ocp-Apim-Subscription-Key': key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(uuid.uuid4())
    }

    body = [{'text': text}]

    try:
        response = requests.post(constructed_url, params=params, headers=headers, json=body)
        response.raise_for_status()
        data = response.json()
        if data and 'translations' in data[0] and data[0]['translations']:
            translated_text = data[0]['translations'][0]['text']
            # Calculation of delay based on the length of the translated text
            delay = len(translated_text) / 10000  # 10 000 characters per second, plan limit Free
            delay = max(1, delay)
            print(f"I wait {delay:.2f} sec. before next API request")
            time.sleep(delay)
            return translated_text
        else:
            print("No translation data in the API response")
            return None
    except requests.exceptions.HTTPError as e:
        print(f"Error HTTP: {e}")
        if response.status_code == 429:
            print("I received error 429, I am waiting 60 seconds")
            time.sleep(60)
        return None
    except requests.exceptions.RequestException as e:
        print(f"Request error: {e}")
    except KeyError as e:
        print(f"Error in response structure: {e}, response: {response.text}")
    except Exception as e:
        print(f"Unidentified error: {e}")
    return None


def read_word_file(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return " ".join(full_text)


def save_text_to_word(text, output_directory, output_file):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.add_run(text)
    
    doc.save(os.path.join(output_directory, output_file))
    print(f"\nTranscription saved to: {os.path.join(output_directory, output_file)}")


source_folder = 'Transcripts'
destination_folder = 'Transcripts-PL'

os.makedirs(destination_folder, exist_ok=True)
for file_name in os.listdir(source_folder):
    if file_name.endswith('.docx'):
        input_file_path = os.path.join(source_folder, file_name)
        output_file_name = file_name.replace('.docx', ' PL.docx')
        output_file_path = os.path.join(destination_folder, output_file_name)

        # Checking whether the file already exists
        if not os.path.exists(output_file_path):
            text_to_translate = read_word_file(input_file_path)
            translated_text = translate_text(text_to_translate)
            if translated_text:
                save_text_to_word(translated_text, destination_folder, output_file_name)
            time.sleep(60)
        else:
            print(f"The file {output_file_name} already exists, I omit the translation")
