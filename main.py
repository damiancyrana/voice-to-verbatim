import os
import time
import azure.cognitiveservices.speech as speechsdk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def find_wav_files(directory):
    """
    Find all WAV files in a given directory in the 'Audio' folder
    """
    audio_dir = os.path.join(directory, 'Audio')
    wav_files = []
    if os.path.exists(audio_dir):
        for file in os.listdir(audio_dir):
            if file.endswith(".wav"):
                wav_files.append(os.path.join(audio_dir, file))
    return wav_files

def recognize_continuous_from_wav(file_path):
    """
    Recognition of speech from a WAV file in continuous mode
    """
    speech_config = speechsdk.SpeechConfig(subscription="", region="")
    audio_config = speechsdk.audio.AudioConfig(filename=file_path)
    speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config, audio_config=audio_config)

    done = False
    all_results = []

    def stop_cb(evt):
        """
        Stop recognition per session event
        """
        print('\nRecognition time completed: {}'.format(evt))
        speech_recognizer.stop_continuous_recognition()
        nonlocal done
        done = True

    def recognized_cb(evt):
        """
        Responds to speech recognition by visualising progress
        """
        if evt.result.text:
            print('.', end='')
            all_results.append(evt.result.text)

    # Connecting events
    speech_recognizer.recognized.connect(recognized_cb)
    speech_recognizer.session_stopped.connect(stop_cb)
    speech_recognizer.canceled.connect(stop_cb)

    # Start recognition
    print("I am starting to recognise...")
    speech_recognizer.start_continuous_recognition()
    while not done:
        time.sleep(0.5)  # Waiting for recognition to complete
    return " ".join(all_results)


def save_text_to_word(text, output_directory, output_file):
    """
    Save the transcription to a Word file with formatting
    """
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.add_run(text)

    doc.save(os.path.join(output_directory, output_file))
    print(f"\nTranscription saved to: {os.path.join(output_directory, output_file)}")


# Set up or create the ‘Transcripts’ folder
current_directory = os.getcwd()
transcripts_directory = os.path.join(current_directory, 'Transcripts')
if not os.path.exists(transcripts_directory):
    os.makedirs(transcripts_directory)

# Finding WAV files and speech recognition
wav_files = find_wav_files(current_directory)
if wav_files:
    for wav_file in wav_files:
        print("Processing a file:", wav_file)
        recognized_text = recognize_continuous_from_wav(wav_file)
        # Creates a Word file name based on the WAV file
        output_file_name = os.path.basename(wav_file).replace('.wav', '.docx')
        if recognized_text:
            save_text_to_word(recognized_text, transcripts_directory, output_file_name)
        else:
            print("Failed to process the file:", wav_file)
else:
    print("WAV files not found in 'Audio' folder")
