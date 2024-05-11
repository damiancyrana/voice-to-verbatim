import subprocess
import json
import openai
from docx import Document
import os
from azure.identity import DefaultAzureCredential
from azure.keyvault.secrets import SecretClient


# Azure Key Vault configuration
key_vault_name = "main-project-key-vault"
key_vault_URI = f"https://{key_vault_name}.vault.azure.net/"
secret_name = "api-key-openAI"

credential = DefaultAzureCredential()
secret_client = SecretClient(vault_url=key_vault_URI, credential=credential)
api_key = secret_client.get_secret(secret_name).value

# OpenAI API configuration
client = openai.OpenAI(
    api_key=api_key
)

TRANSCRIPTS_FOLDER = 'Transcripts-PL'
SUMMARIES_FOLDER = 'Transcript-Summary-PL'
SUMMARIES_LANGUAGE = 'polish'  # polish, english, etc.


def read_word_document(file_path):
    doc = Document(file_path)
    return '\n'.join(paragraph.text for paragraph in doc.paragraphs)


def write_word_document(file_path, text):
    doc = Document()
    for paragraph in text.split('\n'):
        doc.add_paragraph(paragraph)
    doc.save(file_path)


def chat_gpt(prompt):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()


def summarize_text(text, length_percentage=10, language='polish'):
    prompt = f"""You are an assistant specialized in summarizing complex texts concisely and clearly. Summarize the following text to about {length_percentage} % of its original length. The original text is in {language}, and I expect the response in {language}. The Audience are researchers or professionals needing a quick understanding of the topic. Focus on the main points, methodologies, conclusions, and implications. Eliminate redundancies and non-essential details. Text to Summarize: {text}"""
    summary = chat_gpt(prompt)
    return summary


def list_word_documents(folder_path):
    return [f for f in os.listdir(folder_path) if f.endswith('.docx')]


def select_document(documents):
    for index, doc in enumerate(documents, start=1):
        print(f"{index}. {doc}")
    choice = int(input("\nChoose the document number to summarize: "))
    return documents[choice - 1]


def ensure_folder_exists(folder_path):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)


def main():
    transcripts_folder = os.path.join(os.path.dirname(__file__), TRANSCRIPTS_FOLDER)
    summaries_folder = os.path.join(os.path.dirname(__file__), SUMMARIES_FOLDER)
    ensure_folder_exists(summaries_folder)

    documents = list_word_documents(transcripts_folder)
    if not documents:
        print("No Word files to process")
        return

    print("\nAvailable Word files in 'Transcripts-PL':")
    selected_document = select_document(documents)
    file_path = os.path.join(transcripts_folder, selected_document)

    print(f"\nReading file: {file_path}...")
    text = read_word_document(file_path)
    print(f"File loaded: {selected_document}")

    print("Generating summary...")
    summary = summarize_text(text, length_percentage=10, language=SUMMARIES_LANGUAGE)

    base_name, _ = os.path.splitext(selected_document)
    summary_file_path = os.path.join(summaries_folder, f"{base_name}-summary.docx")

    write_word_document(summary_file_path, summary)
    print(f"\nSummary saved to file: {summary_file_path}")


if __name__ == '__main__':
    main()
